import os
from pypdf import PdfReader
from docx import Document as DocxDocument

def extract_text_from_file(file_path: str, file_type: str = None) -> str:
    """
    Extracts raw text from a document of type .pdf, .docx, or .txt.
    If file_type is not provided, it will be inferred from the file extension.
    """
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"File not found at path: {file_path}")

    # Infer file type if not provided
    if not file_type:
        _, ext = os.path.splitext(file_path)
        file_type = ext.lower().lstrip(".")
    else:
        file_type = file_type.lower().lstrip(".")

    text = ""
    try:
        if file_type == "pdf":
            reader = PdfReader(file_path)
            pages_text = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    pages_text.append(page_text)
            text = "\n".join(pages_text)
            
        elif file_type in ("docx", "doc"):
            doc = DocxDocument(file_path)
            paragraphs_text = [p.text for p in doc.paragraphs if p.text]
            # Handle tables inside DOCX
            tables_text = []
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join([cell.text.strip() for cell in row.cells if cell.text.strip()])
                    if row_text:
                        tables_text.append(row_text)
            
            text = "\n".join(paragraphs_text + tables_text)
            
        elif file_type in ("txt", "text"):
            with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
                text = f.read()
                
        else:
            raise ValueError(f"Unsupported file type: '{file_type}'. Only .pdf, .docx, and .txt are supported.")
            
    except Exception as e:
        if isinstance(e, ValueError):
            raise e
        raise Exception(f"Failed to extract text from document of type {file_type}: {str(e)}")

    if not text.strip():
        raise ValueError("Document appears to be empty or contains no extractable text.")

    return text
